from docx import Document
from docx.shared import Pt
import random

def main():
    # Create a new document
    doc = Document()
    
    # Add a title
    doc.add_heading('中文法律文本示例', 0)
    
    # Add some introductory text
    doc.add_paragraph('这是一个包含法律条款格式的示例文档。')
    
    # Generate paragraphs with the pattern "第[零一二三四五六七八九十百千万\d]+条"
    chinese_numbers = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    digits = list(range(10))
    
    # Create a mix of Chinese numbers and digits
    for i in range(1, 21):
        # Randomly choose between Chinese numbers and digits
        if random.choice([True, False]):
            # Use Chinese numbers
            if i < 11:
                number = chinese_numbers[i]
            else:
                number = '十' + chinese_numbers[i-10] if i-10 > 0 else '十'
        else:
            # Use digits
            number = str(i)
        
        # Create the article header
        article_header = f'第{number}条'
        
        # Add a paragraph with the pattern
        p = doc.add_paragraph()
        runner = p.add_run(f'{article_header} ')
        runner.bold = True
        
        # Add some sample text for each article
        p.add_run(f'本条款规定了相关的法律责任和义务，适用于所有相关方。根据《某某法》的规定，当事人需要履行相应的责任。')
    
    # Add a few special cases
    special_cases = [
        '第一百条',
        '第二十五条',
        '第三百四十七条',
        '第1000条',
        '第42条',
    ]
    
    for case in special_cases:
        p = doc.add_paragraph()
        runner = p.add_run(f'{case} ')
        runner.bold = True
        p.add_run(f'特殊情况下，本条款可以适用于例外情形。相关方需按照规定履行义务。')
    
    # Save the document
    output_file = 'regex_pattern_example.docx'
    doc.save(output_file)
    print(f'Document saved as {output_file}')

if __name__ == '__main__':
    main() 